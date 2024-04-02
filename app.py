from flask import Flask, render_template, request, send_file, jsonify, redirect, url_for, session
from werkzeug.utils import secure_filename
from models.keyword import extract_keywords_and_save
from models.ranking import rank_resumes, process_resume_link, DB_retrieve
from models.scraping import scrape_and_get_links
import os
from docx import Document
import pandas as pd
import json
import time

app = Flask(__name__)
app.secret_key = 'your_secret_key' 

conversation_tree = {
    "start": ["Hi! How can I help you? A: Fetching resumes B: Document validation", {"a": "resume_fetch", "b": "document_validation"}],
    "resume_fetch": [
        "How would you like to upload the job description? A: Upload B: Copy Paste",
        {"a": "upload_jd", "b": "copy_paste_jd"}
    ],
    "document_validation": [
        "Please select the type of document for validation. A: Visa B: Driver License",
        {"a": "visa", "b": "driver_license"}
    ],
    "upload_jd": ["Please upload the job description and wait while we process your request...", {"c": "resume_fetch", "d": "start"}],  
    "copy_paste_jd": ["Please paste your job description...", {"c": "resume_fetch", "d": "start"}],  
    "visa": ["Please upload the document for visa validation.", {"c": "document_validation", "d": "start"}],  
    "driver_license": ["Please upload the document for driver's license validation.", {"c": "document_validation", "d": "start"}],  
}

class Chatbot:
    def __init__(self, conversation_tree):
        self.conversation_tree = dict(conversation_tree)
        self.current_node = "start"
        self.pasted_text = None 

    def start_conversation(self):
        return self.conversation_tree[self.current_node][0]

    def process_response(self, response):
        if len(response) > 10:
            filepath='Pasted_jd/pasted_text.txt'
            with open(filepath,'w',encoding='utf-8') as file:
                file.write(response)
            return 'Please click the Fetch button to continue'
        
        options = self.conversation_tree[self.current_node][1]
        if response.lower() in options:
            next_node = options[response.lower()]
            if next_node is not None:
                if next_node == "start":
                    self.current_node = "start"
                elif next_node == "c":
                    prev_node = self.conversation_tree[self.current_node][1].get("c")
                    self.current_node = prev_node
                else:
                    self.current_node = next_node
                return self.conversation_tree[self.current_node][0]
            else:
                self.current_node = "start"
                return "Thank you for uploading. Please wait..."
        else:
            return "Invalid response. Please try again."
    
chatbot = Chatbot(conversation_tree)


def save_to_doc(content, file_path):
    document = Document()
    document.add_paragraph(content)
    document.save(file_path)


@app.route('/')
def index():
    return render_template('index.html', initial_message=chatbot.start_conversation())

@app.route('/response', methods=['POST'])
def get_response():
    global response_text
    response_text = request.json['response']
    next_message = chatbot.process_response(response_text)
    return jsonify({'message': next_message})

@app.route('/process', methods=['POST'])
def process_file():
    global keywords_result
    global ranking_result

    if 'file' in request.files:
        file = request.files['file']

        if file.filename != '':
            filename = secure_filename(file.filename)
            file_path = 'uploads/' + filename
            file.save(file_path)
            keywords_result = extract_keywords_and_save(file_path)
            links_result = scrape_and_get_links([keywords_result])
            ranking_result = rank_resumes(file_path, links_result)
            Job_ID= ranking_result[1]
            ranking_result = ranking_result[0].sort_values(by='Similarity', ascending=False)
            ranking_result = ranking_result.round(2)
            response_data = {'status': 'done'}
            return render_template('dashboard.html', ranking_result=ranking_result.to_dict('records'), Job_ID= Job_ID)

    else:
        a=chatbot.process_response(response_text)
        return render_template('dashboard.html', ranking_result=a)
    return render_template('dashboard.html', error='Error processing file')

@app.route('/process_cp', methods=['POST'])
def process_file_from_path():
    global keywords_result
    global ranking_result
    file_path = 'Pasted_jd/pasted_text.txt'
    keywords_result = extract_keywords_and_save(file_path)
    
    links_result = scrape_and_get_links([keywords_result])
    ranking_result = rank_resumes(file_path, links_result)
    Job_ID= ranking_result[1]
    ranking_result = ranking_result[0].sort_values(by='Similarity', ascending=False)
    ranking_result = ranking_result.round(2)
    ranking_result.to_csv("demo.csv")
    response_data = {'status': 'done'}
    return render_template('dashboard.html', ranking_result=ranking_result.to_dict('records'), Job_ID= Job_ID )

@app.route('/DB', methods=['POST', 'GET'])
def DB():
    job_id = request.form['job_id']
    result_DB = DB_retrieve(job_id)
    result_DB = result_DB.to_dict(orient='records')
    return render_template('dashboard.html', ranking_result=result_DB)


@app.route('/download/<path:link>')
def download(link):
    name, content, email = process_resume_link(link)
    
    doc = Document()
    doc.add_paragraph(content)

    safe_filename = secure_filename(link)
    resumes_name = secure_filename(name) if name else "Unknown"

    download_path = os.path.join('downloaded_resumes', f'{resumes_name}_resume.doc')
    
    os.makedirs(os.path.dirname(download_path), exist_ok=True)

    doc.save(download_path)
    return send_file(download_path, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)

